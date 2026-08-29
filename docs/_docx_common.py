# -*- coding: utf-8 -*-
"""通用 Markdown -> docx 渲染工具（供 数据字典表 / 第三方许可清单 等文档使用）。"""
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

EAST_FONT = "微软雅黑"
LATIN_FONT = "Calibri"
CODE_FONT = "Consolas"

BODY_COLOR = RGBColor(0x33, 0x33, 0x33)
MUTED_COLOR = RGBColor(0x66, 0x66, 0x66)
TITLE_COLOR = RGBColor(0x1F, 0x4D, 0x78)
H1_COLOR = RGBColor(0x2E, 0x74, 0xB5)
H2_COLOR = RGBColor(0x2E, 0x74, 0xB5)
H3_COLOR = RGBColor(0x1F, 0x4D, 0x78)
CALLOUT_COLOR = RGBColor(0x1F, 0x4D, 0x78)
CALLOUT_FILL = "EAF3FB"
CALLOUT_BORDER = "1E90FF"
TABLE_HEADER_FILL = "2E74B5"
TABLE_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)

CONTENT_WIDTH_DXA = 9638  # A4 21cm - 左右2cm边距


def set_font(run, size=11, bold=None, italic=None, color=BODY_COLOR, code=False, east=EAST_FONT):
    latin = CODE_FONT if code else LATIN_FONT
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    run.font.color.rgb = color


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+?\*)")


def add_inline(par, text, size=11, bold=False, color=BODY_COLOR):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            _add_plain(par, text[pos:m.start()], size, bold, color)
        token = m.group(0)
        if token.startswith("**"):
            _add_plain(par, token[2:-2], size, True, color)
        elif token.startswith("`"):
            _add_plain(par, token[1:-1], size - 0.5, bold, color, code=True)
        else:
            _add_plain(par, token[1:-1], size, bold, color, italic=True)
        pos = m.end()
    if pos < len(text):
        _add_plain(par, text[pos:], size, bold, color)


def _add_plain(par, text, size, bold, color, code=False, italic=None):
    if not text:
        return
    run = par.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color, code=code)


def shade_paragraph(p, fill):
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def left_border(p, color, sz=14):
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color)
    pbdr.append(left)
    ppr.append(pbdr)


def style_east_asia(style, east=EAST_FONT):
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), east)


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_COLOR
    style_east_asia(normal)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, H1_COLOR, 16, 8),
        ("Heading 2", 13, H2_COLOR, 12, 6),
        ("Heading 3", 12, H3_COLOR, 10, 4),
    ):
        st = doc.styles[name]
        st.font.name = LATIN_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        style_east_asia(st)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = LATIN_FONT
        st.font.size = Pt(11)
        st.font.color.rgb = BODY_COLOR
        style_east_asia(st)
        st.paragraph_format.space_after = Pt(3)
        st.paragraph_format.line_spacing = 1.25


def setup_section(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(0.6)
    sec.footer_distance = Cm(0.6)


def setup_header_footer(doc, header_text):
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run(header_text)
    set_font(run, size=8.5, color=MUTED_COLOR)

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("第 ")
    set_font(r1, size=9, color=MUTED_COLOR)
    fld = fp.add_run()
    set_font(fld, size=9, color=MUTED_COLOR)
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    fld._r.append(f1); fld._r.append(instr); fld._r.append(f2)
    r2 = fp.add_run(" 页")
    set_font(r2, size=9, color=MUTED_COLOR)


def numbering_element(doc):
    return doc.part.numbering_part.element


def style_num_id(doc, style_name):
    style = doc.styles[style_name]
    num_id_el = style.element.find(qn("w:pPr") + "/" + qn("w:numPr") + "/" + qn("w:numId"))
    if num_id_el is None:
        return None
    return int(num_id_el.get(qn("w:val")))


def abstract_num_id(doc, num_id):
    numbering = numbering_element(doc)
    for num in numbering.findall(qn("w:num")):
        if int(num.get(qn("w:numId"))) == num_id:
            ref = num.find(qn("w:abstractNumId"))
            if ref is not None:
                return int(ref.get(qn("w:val")))
    return None


def new_num_id(doc, abstract_id):
    numbering = numbering_element(doc)
    ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    new_id = max(ids) + 1 if ids else 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return new_id


def apply_num(paragraph, num_id, ilvl=0):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    lvl = OxmlElement("w:ilvl"); lvl.set(qn("w:val"), str(ilvl))
    nid = OxmlElement("w:numId"); nid.set(qn("w:val"), str(num_id))
    numpr.append(lvl); numpr.append(nid)
    ppr.append(numpr)


def shade_cell(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def apply_geometry(table, widths_cm):
    total_cm = sum(widths_cm)
    dxas = [int(round(w / total_cm * CONTENT_WIDTH_DXA)) for w in widths_cm]
    dxas[-1] += CONTENT_WIDTH_DXA - sum(dxas)
    tbl = table._tbl
    tblpr = tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW"); tblpr.append(tblw)
    tblw.set(qn("w:w"), str(CONTENT_WIDTH_DXA)); tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd"); tblpr.append(tblind)
    tblind.set(qn("w:w"), "120"); tblind.set(qn("w:type"), "dxa")
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout"); tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.find(qn("w:tblGrid"))
    for col, dxa in zip(grid.findall(qn("w:gridCol")), dxas):
        col.set(qn("w:w"), str(dxa))
    for row in table.rows:
        for cell, dxa in zip(row.cells, dxas):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW"); tcpr.append(tcw)
            tcw.set(qn("w:w"), str(dxa)); tcw.set(qn("w:type"), "dxa")


def auto_widths(rows):
    """按各列最长内容比例分配列宽；2/3 列用固定方案与既有文档一致。"""
    n_cols = len(rows[0])
    if n_cols == 2:
        return [4.2, 12.8]
    if n_cols == 3:
        return [4.6, 3.2, 9.2]
    lengths = []
    for ci in range(n_cols):
        longest = 0
        for r in rows:
            v = r[ci] if ci < len(r) else ""
            # 中文按 1 宽、ASCII 按 0.5 宽估算
            w = sum(1 if ord(ch) > 0x2E7F else 0.5 for ch in v)
            longest = max(longest, w)
        # 单列估算上限 14、下限 5，避免长路径/长文本挤压其他列
        lengths.append(max(min(longest, 14), 5))
    total = sum(lengths)
    widths = [max(1.4, 17.0 * l / total) for l in lengths]
    scale = 17.0 / sum(widths)
    return [w * scale for w in widths]


def make_table(doc, rows):
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, row in enumerate(rows):
        for ci in range(n_cols):
            cell_text = row[ci] if ci < len(row) else ""
            cell = table.cell(ri, ci)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            if ri == 0:
                shade_cell(cell, TABLE_HEADER_FILL)
                add_inline(p, cell_text, size=10.5, bold=True, color=TABLE_HEADER_TEXT)
            else:
                add_inline(p, cell_text, size=10.5)
    apply_geometry(table, auto_widths(rows))
    return table


def add_callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    shade_paragraph(p, CALLOUT_FILL)
    left_border(p, CALLOUT_BORDER)
    add_inline(p, text, size=10.5, color=CALLOUT_COLOR)


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    shade_paragraph(p, "F5F7FA")
    for i, ln in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(ln)
        set_font(run, size=10, color=RGBColor(0x1F, 0x4D, 0x78), code=True)


def parse_md(text):
    lines = text.splitlines()
    blocks = []
    seen_rule = False
    in_code = False
    code_lines = []
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        if line.strip() == "```":
            if in_code:
                blocks.append(("code", code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        if re.match(r"^\s*---+\s*$", line):
            seen_rule = True
            blocks.append(("rule", None))
            i += 1
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            blocks.append(("h", (level, m.group(2).strip())))
            i += 1
            continue
        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            if len(rows) >= 2 and all(re.fullmatch(r":?-{2,}:?", c) for c in rows[1]):
                rows.pop(1)
            blocks.append(("table", rows))
            continue
        if line.startswith(">"):
            blocks.append(("callout", re.sub(r"^>\s?", "", line).strip()))
            i += 1
            continue
        m = re.match(r"^(\d+)\.\s+(.*)$", line)
        if m:
            blocks.append(("num", m.group(2)))
            i += 1
            continue
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            level = 2 if len(m.group(1)) >= 2 else 1
            blocks.append(("bullet", (level, m.group(2))))
            i += 1
            continue
        blocks.append(("para", line.strip()))
        i += 1
    return blocks


def render(md_path, out_path, header_text, core_title, core_author):
    md_path = Path(md_path)
    out_path = Path(out_path)
    text = md_path.read_text(encoding="utf-8-sig")
    blocks = parse_md(text)

    doc = Document()
    setup_styles(doc)
    setup_section(doc)
    setup_header_footer(doc, header_text)
    doc.core_properties.title = core_title
    doc.core_properties.author = core_author

    title_done = False
    in_cover = False
    numbered = {"num_id": None, "open": False}

    def close_list():
        numbered["open"] = False

    for kind, payload in blocks:
        if kind == "h":
            level, t = payload
            if level == 1 and not title_done:
                close_list()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(4)
                add_inline(p, t, 22, bold=True, color=TITLE_COLOR)
                title_done = True
                in_cover = True
                continue
            in_cover = False
            close_list()
            if level == 2:
                p = doc.add_paragraph(style="Heading 1")
                add_inline(p, t, size=16, bold=True, color=H1_COLOR)
            elif level == 3:
                p = doc.add_paragraph(style="Heading 2")
                add_inline(p, t, size=13, bold=True, color=H2_COLOR)
            else:
                p = doc.add_paragraph(style="Heading 3")
                add_inline(p, t, size=12, bold=True, color=H3_COLOR)
            continue
        if kind == "rule":
            in_cover = False
            close_list()
            continue
        if kind == "para":
            close_list()
            p = doc.add_paragraph()
            if in_cover:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(2)
                add_inline(p, payload, size=10.5, color=MUTED_COLOR)
            else:
                add_inline(p, payload, size=11)
            continue
        if kind == "callout":
            in_cover = False
            close_list()
            add_callout(doc, payload)
            continue
        if kind == "code":
            in_cover = False
            close_list()
            add_code_block(doc, payload)
            continue
        if kind == "table":
            in_cover = False
            close_list()
            make_table(doc, payload)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        if kind == "bullet":
            level, text = payload
            style_name = "List Bullet" if level == 1 else "List Bullet 2"
            try:
                p = doc.add_paragraph(style=style_name)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(0.75 if level == 2 else 0.63)
            add_inline(p, text, size=11)
            continue
        if kind == "num":
            if not numbered["open"]:
                base_id = style_num_id(doc, "List Number")
                abstract_id = abstract_num_id(doc, base_id) if base_id is not None else None
                if abstract_id is None:
                    abstract_id = 0
                numbered["num_id"] = new_num_id(doc, abstract_id)
                numbered["open"] = True
            p = doc.add_paragraph(style="List Number")
            apply_num(p, numbered["num_id"])
            add_inline(p, payload, size=11)
            continue

    doc.save(out_path)
    print(f"[OK] {out_path}")
